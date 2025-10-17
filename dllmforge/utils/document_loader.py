"""
Document loader module for extracting text from various document formats.
Each document type handler is optional - if the required package is not installed,
the corresponding functionality will be disabled but other formats will still work.
"""
from pathlib import Path
from typing import Union, Optional
import importlib.util


class DocumentLoader:
    """Class for loading and extracting text from various document formats"""

    def __init__(self):
        """Initialize document loader and check available handlers"""
        self.handlers = {
            '.pdf': self._extract_from_pdf if self._check_pdf_support() else None,
            '.docx': self._extract_from_word if self._check_docx_support() else None,
            '.xlsx': self._extract_from_excel if self._check_excel_support() else None,
            '.xls': self._extract_from_excel if self._check_excel_support() else None,
            '.csv': self._extract_from_csv if self._check_csv_support() else None
        }

        # Print available handlers
        available = [ext for ext, handler in self.handlers.items() if handler is not None]
        print(f"Document loader initialized with support for: {', '.join(available)}")

    @staticmethod
    def _check_pdf_support() -> bool:
        """Check if PDF support is available"""
        return (importlib.util.find_spec('PyPDF2') is not None and importlib.util.find_spec('pdf2image') is not None
                and importlib.util.find_spec('pytesseract') is not None)

    @staticmethod
    def _check_docx_support() -> bool:
        """Check if Word document support is available"""
        return importlib.util.find_spec('docx') is not None

    @staticmethod
    def _check_excel_support() -> bool:
        """Check if Excel support is available"""
        return importlib.util.find_spec('pandas') is not None

    @staticmethod
    def _check_csv_support() -> bool:
        """Check if CSV support is available"""
        return importlib.util.find_spec('pandas') is not None

    def load_document(self, file_path: Union[str, Path]) -> str:
        """Load document from various formats and return text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is unsupported or required package is not installed
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        handler = self.handlers.get(suffix)
        if handler is None:
            if suffix in self.handlers:
                raise ValueError(f"Required packages for {suffix} files are not installed. "
                                 f"Please install the necessary packages.")
            else:
                raise ValueError(f"Unsupported file format: {suffix}")

        return handler(file_path)

    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        from PyPDF2 import PdfReader
        from pdf2image import convert_from_path
        import pytesseract

        # First try direct PDF text extraction
        pdf = PdfReader(str(file_path))
        text = ""
        for page in pdf.pages:
            text += page.extract_text()

        # If no text was extracted, use OCR
        if not text.strip():
            images = convert_from_path(str(file_path))
            text = ""
            for image in images:
                text += pytesseract.image_to_string(image)

        return text

    def _extract_from_word(self, file_path: Path) -> str:
        """Extract text from Word documents"""
        from docx import Document

        doc = Document(str(file_path))
        text = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text.append(" | ".join(row_text))

        return "\n".join(text)

    def _extract_from_excel(self, file_path: Path) -> str:
        """Extract text from Excel files"""
        import pandas as pd
        df = pd.read_excel(str(file_path))
        return df.to_string()

    def _extract_from_csv(self, file_path: Path) -> str:
        """Extract text from CSV files"""
        import pandas as pd
        df = pd.read_csv(str(file_path))
        return df.to_string()
